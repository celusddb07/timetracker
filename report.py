from datetime import timedelta
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _set_cell_bg(cell, hex_color: str):
    """Set a table cell background colour (hex without #)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def build_report(week_start, week_end, entries):
    """Return a BytesIO containing the .docx report."""
    doc = Document()

    # ---- Title ----
    title = doc.add_heading(level=0)
    title.clear()
    run = title.add_run(
        f"Time Report – Week of "
        f"{week_start.strftime('%d %b')} – {week_end.strftime('%d %b %Y')}"
    )
    run.bold = True

    # ---- Group entries by day ----
    days = {}
    for i in range(7):
        days[week_start + timedelta(days=i)] = []
    for entry in entries:
        days[entry["entry_date"]].append(entry)

    grand_total = 0.0

    for day, day_entries in days.items():
        if not day_entries:
            continue

        # Day heading
        doc.add_heading(day.strftime("%A, %d %B %Y"), level=2)

        # Table with header row
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Hours"
        hdr_cells[1].text = "Description"
        hdr_cells[2].text = "Concepts Learned"
        for cell in hdr_cells:
            _set_cell_bg(cell, "D9D9D9")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        day_total = 0.0
        for entry in day_entries:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{entry['hours']:.2f}"
            row_cells[1].text = entry["description"]
            row_cells[2].text = entry["concepts_learned"] or ""
            day_total += entry["hours"]

        # Day subtotal row
        sub_cells = table.add_row().cells
        sub_cells[0].text = f"{day_total:.2f}"
        sub_cells[1].text = "Day total"
        sub_cells[2].text = ""
        _set_cell_bg(sub_cells[0], "EBF3FB")
        _set_cell_bg(sub_cells[1], "EBF3FB")
        _set_cell_bg(sub_cells[2], "EBF3FB")
        for cell in sub_cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True

        grand_total += day_total
        doc.add_paragraph("")  # spacing after table

    # ---- Grand total ----
    p = doc.add_paragraph()
    run = p.add_run(f"Total hours this week: {grand_total:.2f}")
    run.bold = True
    run.font.size = Pt(12)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
